
import streamlit as st
import os
import tempfile
import base64
from langchain_community.document_loaders import PyPDFLoader, Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_groq import ChatGroq
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from docx import Document
from weasyprint import HTML
import markdown
import io
import uuid
import shutil
from PIL import Image
import PyPDF2
import fitz  # PyMuPDF
import re

# Set page configuration
st.set_page_config(page_title="AI Proposal Generator", layout="wide")

# Initialize session state variables if they don't exist
if "knowledge_base" not in st.session_state:
    st.session_state.knowledge_base = None
if "knowledge_files" not in st.session_state:
    st.session_state.knowledge_files = []
if "extracted_images" not in st.session_state:
    st.session_state.extracted_images = []
if "proposal_template" not in st.session_state:
    st.session_state.proposal_template = """
# Proposal

## Introduction
[Brief introduction about the company]

## Scope/Objectives
[Define the scope and objectives of the project]

## Proposal/Approach
[Detail the approach to meet the objectives]

## Proposed Workflow
```mermaid
graph TD
    A[Requirement Analysis] --> B[Design]
    B --> C[Development]
    C --> D[Testing]
    D --> E[Deployment]
    E --> F[Support]
```

## Deliverables from Client
[List what is expected from the client]

## Timelines
[Project timeline with milestones]

## Commercials
[Cost breakdown and payment terms]

## About AI Planet
[Information about AI Planet]
"""
if "requirements" not in st.session_state:
    st.session_state.requirements = ""
if "generated_proposal" not in st.session_state:
    st.session_state.generated_proposal = ""
if "template_option" not in st.session_state:
    st.session_state.template_option = "Use Default"
if "section_generation" not in st.session_state:
    st.session_state.section_generation = {
        "current_section": 0,
        "sections": [],
        "generated_sections": []
    }

# Create a directory for storing images
IMAGE_DIR = "extracted_images"
if not os.path.exists(IMAGE_DIR):
    os.makedirs(IMAGE_DIR)

# Helper functions
def extract_text_and_images_from_pdf(file_path):
    """Extract text and images from PDF file using LangChain's PyPDFLoader"""
    extracted_text = ""
    image_paths = []

    try:
        # Use LangChain's PyPDFLoader for better text extraction
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        extracted_text = "\n\n".join([doc.page_content for doc in docs])

        # Extract images using PyMuPDF
        try:
            # Open the PDF with PyMuPDF
            pdf_document = fitz.Document(file_path)

            # Iterate through pages
            for page_num, page in enumerate(pdf_document):
                # Extract images
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    base_image = pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]

                    # Generate a unique filename
                    image_filename = f"{IMAGE_DIR}/pdf_image_{uuid.uuid4().hex}.png"

                    # Save the image
                    with open(image_filename, "wb") as img_file:
                        img_file.write(image_bytes)

                    image_paths.append(image_filename)

            pdf_document.close()
        except Exception as e:
            st.warning(f"Image extraction failed: {str(e)}. Only text will be extracted.")

    except Exception as e:
        st.error(f"Error extracting from PDF: {str(e)}")

    return extracted_text, image_paths

def extract_text_and_images_from_docx(file_path):
    """Extract text and images from DOCX file using LangChain's Docx2txtLoader"""
    # Extract text using LangChain's loader
    loader = Docx2txtLoader(file_path)
    docs = loader.load()
    extracted_text = "\n\n".join([doc.page_content for doc in docs])

    # Extract images
    image_paths = []
    try:
        doc = Document(file_path)

        # Extract all the shapes (including images)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                # Save the image
                image_filename = f"{IMAGE_DIR}/docx_image_{uuid.uuid4().hex}.png"
                with open(image_filename, "wb") as img_file:
                    img_file.write(rel.target_part.blob)
                image_paths.append(image_filename)
    except Exception as e:
        st.warning(f"Image extraction from DOCX failed: {str(e)}. Only text will be extracted.")

    return extracted_text, image_paths

def extract_text_from_file(uploaded_file):
    """Extract text from uploaded file based on file type"""
    # Create a temporary file to store the uploaded content
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
        tmp_file.write(uploaded_file.getbuffer())
        file_path = tmp_file.name

    try:
        # Use the appropriate loader based on file type
        if file_path.endswith(".pdf"):
            extracted_text, image_paths = extract_text_and_images_from_pdf(file_path)
            st.session_state.extracted_images.extend(image_paths)
        elif file_path.endswith(".docx"):
            extracted_text, image_paths = extract_text_and_images_from_docx(file_path)
            st.session_state.extracted_images.extend(image_paths)
        else:
            raise ValueError("Only PDF and DOCX files are supported")

        # Clean up the temporary file
        os.unlink(file_path)
        return extracted_text
    except Exception as e:
        os.unlink(file_path)  # Clean up on error
        raise e

def create_kb_from_texts(texts):
    """Create knowledge base from list of texts"""
    combined_text = "\n\n".join(texts)
    # Create text chunks for the vector store
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    chunks = text_splitter.split_text(combined_text)

    # Create embeddings using Hugging Face
    embeddings = HuggingFaceEmbeddings(model_name=selected_embedding_model)
    return FAISS.from_texts(chunks, embeddings)

def convert_md_to_docx(md_text, output_filename, images=None):
    """Convert markdown to DOCX with images"""
    doc = Document()
    # Add any available logos at the top
    if images:
        for img_path in images[:2]:  # Limit to first 2 images which are likely logos
            try:
                doc.add_picture(img_path, width=2000000)  # ~2 inches wide
            except Exception as e:
                pass  # Skip if image can't be added

    # Simple parsing of markdown headings and paragraphs
    lines = md_text.split('\n')
    
    # Track if we're in a mermaid code block
    in_mermaid = False
    mermaid_content = []
    
    for line in lines:
        line = line.strip()
        if line.startswith("```mermaid"):
            in_mermaid = True
            continue
        elif in_mermaid and line.startswith("```"):
            in_mermaid = False
            # Add a placeholder for the mermaid diagram
            doc.add_paragraph("--- Mermaid Workflow Diagram ---")
            continue
        elif in_mermaid:
            mermaid_content.append(line)
            continue
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('```') and line != '```':
            # Skip code blocks
            continue
        elif '```' in line:
            # Skip code blocks
            continue
        elif line:
            doc.add_paragraph(line)

    # Save the document
    doc.save(output_filename)
    return output_filename

def convert_md_to_pdf(md_text, output_filename, images=None):
    """Convert markdown to PDF with images"""
    # Process markdown text and handle mermaid diagrams
    processed_md = md_text
    
    # Convert markdown to HTML (excluding mermaid blocks which will be handled separately)
    html = markdown.markdown(processed_md, extensions=['extra', 'codehilite'])
    
    # Add image tags if images are available
    image_tags = ""
    if images:
        for img_path in images[:2]:  # Limit to first 2 images which are likely logos
            try:
                # Convert to base64 for embedding in HTML
                with open(img_path, "rb") as img_file:
                    img_data = base64.b64encode(img_file.read()).decode()
                    img_ext = os.path.splitext(img_path)[1][1:]
                    image_tags += f'<img src="data:image/{img_ext};base64,{img_data}" style="max-width: 200px; margin: 10px 0;">'
            except Exception as e:
                pass  # Skip if image can't be processed

    # Replace mermaid sections with placeholders in HTML
    html = re.sub(r'<pre><code>mermaid.*?</code></pre>', '<div class="mermaid-placeholder">--- Mermaid Workflow Diagram ---</div>', html, flags=re.DOTALL)
    
    styled_html = f"""
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 2cm; }}
            h1 {{ color: #2c3e50; }}
            h2 {{ color: #3498db; }}
            h3 {{ color: #555; }}
            code {{ background-color: #f8f8f8; padding: 2px 4px; }}
            .mermaid-placeholder {{ 
                text-align: center;
                padding: 20px;
                border: 1px dashed #ccc;
                margin: 20px 0;
                color: #777;
                font-style: italic;
            }}
        </style>
    </head>
    <body>
        {image_tags}
        {html}
    </body>
    </html>
    """
    HTML(string=styled_html).write_pdf(output_filename)
    return output_filename

def parse_template_sections(template):
    """Parse the proposal template into sections"""
    sections = []
    lines = template.strip().split('\n')
    
    current_section = None
    
    for line in lines:
        if line.startswith('## '):
            # If we have content in the current section, save it
            if current_section:
                sections.append(current_section)
            # Start a new sub-section
            current_section = {"title": line[3:], "content": line + "\n", "lines": [line]}
        elif current_section:
            # Add line to current section
            current_section["content"] += line + "\n"
            current_section["lines"].append(line)
    
    # Add the last section
    if current_section:
        sections.append(current_section)
        
    return sections

def generate_custom_workflow_mermaid(requirements, llm):
    """Generate a custom mermaid workflow diagram based on project requirements"""
    prompt_template = """
    You are a project workflow specialist. Based on the project requirements below, create a custom project workflow diagram using mermaid syntax. 
    
    # PROJECT REQUIREMENTS:
    {requirements}
    
    Instructions:
    1. Analyze the project requirements carefully
    2. Create a detailed, logical workflow that would be appropriate for executing this specific project
    3. Include key phases, deliverables, and dependencies
    4. Use appropriate node shapes and connections
    5. The workflow should be specific to this project, not generic
    6. Use TD (top-down) orientation
    7. Name each node with a letter and short descriptive text (e.g., A[Requirement Analysis])
    8. Return ONLY valid mermaid graph syntax for a workflow diagram, nothing else
    9. Use this format:
    
    graph TD
        A[First Step] --> B[Second Step]
        B --> C[Third Step]
        ...etc
    
    The workflow must be specifically tailored to the requirements provided and reflect a realistic project execution approach.
    """
    
    # Create the prompt
    prompt = ChatPromptTemplate.from_template(prompt_template)
    chain = prompt | llm
    
    # Generate the mermaid workflow diagram
    result = chain.invoke({"requirements": requirements})
    
    # Clean up the result to ensure it has the right format
    result_content = result.content.strip()
    
    # Ensure the result is valid mermaid syntax
    if not result_content.startswith("graph TD"):
        result_content = "graph TD\n" + result_content
        
    return result_content

def generate_section_content(section, requirements, context, llm):
    """Generate content for a specific section of the proposal"""
    section_title = section["title"]
    section_content = section["content"]
    
    # Special handling for the workflow section
    if "Workflow" in section_title:
        try:
            # Generate a custom mermaid workflow based on requirements
            custom_workflow = generate_custom_workflow_mermaid(requirements, llm)
            
            # Format the section with the custom workflow
            workflow_section = f"## {section_title}\n\n```mermaid\n{custom_workflow}\n```"
            return workflow_section
        except Exception as e:
            st.warning(f"Error generating custom workflow: {str(e)}. Using default workflow.")
            # Return the original section content if there's an error
            return section_content
    
    # For other sections, use the standard generation approach
    # Create a prompt to generate content specifically for this section
    prompt_template = f"""
    You are a professional proposal writer tasked with creating a detailed section for a business proposal.

    # SECTION TO COMPLETE
    {section_content}

    # PROJECT REQUIREMENTS
    {requirements}

    # RELEVANT CONTEXT
    {context}

    Instructions:
    1. Generate content ONLY for the specific section titled "{section_title}".
    2. Keep the section header format exactly as shown above.
    3. Replace any placeholder text in brackets with detailed, specific content.
    4. Use the context information but DO NOT directly copy company names or specific identifiers from the context.
    5. Maintain a professional, business proposal tone.
    6. Be detailed and thorough but concise.
    7. If this is a technical section, provide specific technical details.
    8. If this section includes diagrams or charts, maintain those placeholders.

    IMPORTANT: Generate ONLY the content for this specific section, nothing more.
    """
    
    # Create the prompt
    prompt = ChatPromptTemplate.from_template(prompt_template)
    chain = prompt | llm
    
    # Generate the section content
    result = chain.invoke({})
    
    # Clean up the result to ensure it has the right format
    result_content = result.content.strip()
    
    # Make sure we preserve section headers
    if section_content.startswith('# ') and not result_content.startswith('# '):
        result_content = section_content.split('\n')[0] + '\n\n' + result_content
    elif section_content.startswith('## ') and not result_content.startswith('## '):
        result_content = section_content.split('\n')[0] + '\n\n' + result_content
        
    return result_content

def get_complete_proposal():
    """Combine all generated sections into a complete proposal"""
    # If no sections have been generated yet, return an empty string
    if not st.session_state.section_generation["generated_sections"]:
        return ""
    
    # Start with the main proposal title
    complete_proposal = "# Proposal\n\n"
    
    # Add all generated sections
    for section in st.session_state.section_generation["generated_sections"]:
        complete_proposal += section + "\n\n"
        
    return complete_proposal

# Set up sidebar with API keys and model selection
st.sidebar.title("Configuration")

# Set up API key input for Groq
groq_api_key = st.sidebar.text_input("Enter Groq API Key", type="password")
os.environ["GROQ_API_KEY"] = groq_api_key

# Set up embedding model selection
embedding_model_options = [
    "sentence-transformers/all-MiniLM-L6-v2",  # Fast and lightweight
    "sentence-transformers/all-mpnet-base-v2",  # Better quality
    "BAAI/bge-small-en-v1.5"  # Another good option
]
selected_embedding_model = st.sidebar.selectbox(
    "Select Embedding Model",
    embedding_model_options,
    index=0
)

# Groq LLM model selection
groq_model_options = [
    "llama3-70b-8192",
    "llama3-8b-8192",
    "mixtral-8x7b-32768",
    "gemma-7b-it",
    "claude-3-opus-20240229",
    "claude-3-sonnet-20240229"
]
selected_llm_model = st.sidebar.selectbox(
    "Select Groq LLM Model",
    groq_model_options,
    index=0
)

# Main UI
st.title("🤖 AI Proposal Generator")

# Create columns for main layout
col1, col2 = st.columns([1, 1])

with col1:
    # Requirements section
    st.header("1️⃣ Project Requirements")

    # Input methods
    input_method = st.radio("Choose input method:", ["Text Input", "File Upload"], horizontal=True, key="req_input_method")

    if input_method == "Text Input":
        new_requirements = st.text_area(
            "Enter the project requirements:",
            height=150,
            value=st.session_state.requirements
        )
        # Update session state when text changes
        if new_requirements != st.session_state.requirements:
            st.session_state.requirements = new_requirements
    else:
        uploaded_file = st.file_uploader("Upload requirements document", type=["pdf", "docx"], key="req_file")
        if uploaded_file:
            try:
                st.session_state.requirements = extract_text_from_file(uploaded_file)
                st.success(f"Successfully extracted requirements from {uploaded_file.name}")

                # Show the extracted text
                with st.expander("View Extracted Requirements"):
                    st.text_area("Extracted content:", value=st.session_state.requirements, height=150)
            except Exception as e:
                st.error(f"Error extracting text: {str(e)}")

    # Knowledge base section
    st.header("2️⃣ Knowledge Base")

    # Allow uploading multiple files
    kb_files = st.file_uploader("Upload knowledge base documents", type=["pdf", "docx"], accept_multiple_files=True, key="kb_files")

    # Add files to knowledge base list
    if kb_files:
        new_files = [f for f in kb_files if f.name not in [kf.name for kf in st.session_state.knowledge_files]]
        st.session_state.knowledge_files.extend(new_files)

    # Display existing knowledge base files
    if st.session_state.knowledge_files:
        st.write("Knowledge Base Files:")
        cols = st.columns(3)
        for i, file in enumerate(st.session_state.knowledge_files):
            col_idx = i % 3
            with cols[col_idx]:
                st.write(f"📄 {file.name}")

    # Option to clear knowledge base
    if st.session_state.knowledge_files and st.button("Clear Knowledge Base"):
        st.session_state.knowledge_files = []
        st.session_state.knowledge_base = None
        st.session_state.extracted_images = []
        # Clean up image directory
        for filename in os.listdir(IMAGE_DIR):
            file_path = os.path.join(IMAGE_DIR, filename)
            try:
                if os.path.isfile(file_path):
                    os.unlink(file_path)
            except Exception as e:
                st.warning(f"Error deleting {file_path}: {e}")
        st.success("Knowledge base cleared.")

    # Create knowledge base button
    if st.button("Create Knowledge Base"):
        if not st.session_state.knowledge_files and not st.session_state.requirements:
            st.warning("Please add requirements or upload knowledge base files first.")
        else:
            # Show progress message
            with st.spinner(f"Creating Knowledge Base"):
                try:
                    # Extract text from all knowledge base files
                    kb_texts = []

                    # Add requirements if available
                    if st.session_state.requirements:
                        kb_texts.append(st.session_state.requirements)

                    # Add content from all files
                    for file in st.session_state.knowledge_files:
                        try:
                            file_content = extract_text_from_file(file)
                            kb_texts.append(file_content)
                        except Exception as e:
                            st.warning(f"Could not process {file.name}: {str(e)}")

                    # Create the vector database knowledge base
                    st.session_state.knowledge_base = create_kb_from_texts(kb_texts)
                    st.success("✅ Knowledge base created successfully!")
                except Exception as e:
                    st.error(f"Error creating knowledge base: {str(e)}")

    # Proposal template section
    st.header("3️⃣ Proposal Template")

    template_options = ["Use Default", "Edit Template", "Upload Template"]
    st.session_state.template_option = st.radio("Choose template method:", template_options, horizontal=True, key="template_method")

    if st.session_state.template_option == "Use Default":
        with st.expander("View Default Template"):
            st.code(st.session_state.proposal_template, language="markdown")
    elif st.session_state.template_option == "Edit Template":
        new_template = st.text_area(
            "Edit the proposal template below:",
            value=st.session_state.proposal_template,
            height=250
        )
        # Update session state when template changes
        if new_template != st.session_state.proposal_template:
            st.session_state.proposal_template = new_template
            # Reset section generation when template changes
            st.session_state.section_generation = {
                "current_section": 0,
                "sections": parse_template_sections(new_template),
                "generated_sections": []
            }
    else:  # Upload Template
        uploaded_template = st.file_uploader("Upload proposal template", type=["docx", "pdf"], key="template_file")
        if uploaded_template:
            try:
                content = extract_text_from_file(uploaded_template)
                st.session_state.proposal_template = content
                # Reset section generation when template changes
                st.session_state.section_generation = {
                    "current_section": 0,
                    "sections": parse_template_sections(content),
                    "generated_sections": []
                }
                with st.expander("Template Preview"):
                    st.text_area("Template Content:", value=content, height=200)
            except Exception as e:
                st.error(f"Error extracting template: {str(e)}")

with col2:
    # Generation and results section
    st.header("4️⃣ Generate & Edit Proposal")

    # Parse template into sections if needed
    if (st.session_state.section_generation["sections"] == [] and 
        st.session_state.proposal_template):
        st.session_state.section_generation["sections"] = parse_template_sections(
            st.session_state.proposal_template
        )
        st.session_state.section_generation["current_section"] = 0
        st.session_state.section_generation["generated_sections"] = []

    # Show which section is being generated next
    if st.session_state.section_generation["sections"]:
        current_idx = st.session_state.section_generation["current_section"]
        total_sections = len(st.session_state.section_generation["sections"])
        
        if current_idx < total_sections:
            next_section = st.session_state.section_generation["sections"][current_idx]["title"]
            st.info(f"Next section to generate: **{next_section}**")
        else:
            st.success("All sections have been generated!")

    # Generate proposal button - now generates section by section
    if st.button("🚀 Generate Next Section", use_container_width=True):
        if not groq_api_key:
            st.warning("Please enter your Groq API key in the sidebar.")
        elif not st.session_state.requirements:
            st.warning("Please enter project requirements first.")
        elif not st.session_state.knowledge_base and st.session_state.knowledge_files:
            st.warning("Please create a knowledge base first.")
        else:
            try:
                # Get current section index
                current_idx = st.session_state.section_generation["current_section"]
                
                # Check if we have sections to generate
                if current_idx >= len(st.session_state.section_generation["sections"]):
                    st.warning("All sections have been generated!")
                else:
                    # Get the current section to generate
                    current_section = st.session_state.section_generation["sections"][current_idx]
                    
                    with st.spinner(f"Generating '{current_section['title']}' section "):
                        # Set up the Groq LLM
                        llm = ChatGroq(
                            model=selected_llm_model,
                            temperature=0.2,
                            api_key=groq_api_key
                        )

                        # Get the requirements from session state
                        requirements = st.session_state.requirements

                        # Get relevant context from knowledge base if it exists
                        context = ""
                        if st.session_state.knowledge_base:
                            retriever = st.session_state.knowledge_base.as_retriever(
                                search_type="similarity",
                                search_kwargs={"k": 5}
                            )
                            # Get relevant context specifically for this section
                            section_query = f"{current_section['title']} {requirements}"
                            context_docs = retriever.get_relevant_documents(section_query)
                            context = "\n\n".join([doc.page_content for doc in context_docs])
                        
                        # Generate content for this section
                        section_content = generate_section_content(
                            current_section, 
                            requirements, 
                            context, 
                            llm
                        )
                        
                        # Add to generated sections
                        st.session_state.section_generation["generated_sections"].append(section_content)
                        
                        # Advance to next section
                        st.session_state.section_generation["current_section"] += 1
                        
                        # Update the current section in the editing area
                        st.session_state.generated_proposal = section_content
                        
                        st.success(f"Section '{current_section['title']}' generated successfully!")

            except Exception as e:
                st.error(f"Error generating section: {str(e)}")

    # Show progress of section generation
    if st.session_state.section_generation["sections"]:
        total_sections = len(st.session_state.section_generation["sections"])
        current_section = st.session_state.section_generation["current_section"]
        
        st.progress(min(1.0, current_section / total_sections), 
                   text=f"Sections generated: {current_section}/{total_sections}")
        
        # Option to restart section generation
        if current_section > 0 and st.button("Restart Section Generation"):
            st.session_state.section_generation["current_section"] = 0
            st.session_state.section_generation["generated_sections"] = []
            st.session_state.generated_proposal = ""
            st.success("Section generation reset. You can generate sections again.")

    # Display and edit the most recently generated section
    if st.session_state.generated_proposal:
        # Show only the most recently generated section for editing
        st.subheader("Edit Current Section")
        edited_section = st.text_area(
            "Edit this section as needed:",
            value=st.session_state.generated_proposal,
            height=300
        )
        
        # Update the section in session state if it was edited
        if edited_section != st.session_state.generated_proposal:
            # Find the index of the most recently generated section
            if st.session_state.section_generation["current_section"] > 0:
                section_idx = st.session_state.section_generation["current_section"] - 1
                # Update the section content
                st.session_state.section_generation["generated_sections"][section_idx] = edited_section
            st.session_state.generated_proposal = edited_section

        # Display complete proposal preview
        with st.expander("Proposal Preview (All Sections)"):
            # Get the complete proposal with all sections
            complete_proposal = get_complete_proposal()
            st.markdown(complete_proposal)

        # Create temporary files for download
        with tempfile.NamedTemporaryFile(delete=False, suffix=".md") as md_file:
            complete_proposal = get_complete_proposal()
            md_file.write(complete_proposal.encode())
            md_path = md_file.name

        docx_path = md_path.replace(".md", ".docx")
        pdf_path = md_path.replace(".md", ".pdf")

        try:
            # Convert to DOCX and PDF with extracted images
            convert_md_to_docx(complete_proposal, docx_path, st.session_state.extracted_images)
            convert_md_to_pdf(complete_proposal, pdf_path, st.session_state.extracted_images)

            # Create download buttons
            col1, col2 = st.columns(2)

            with col1:
                st.download_button(
                    label="Download as DOCX",
                    data=open(docx_path, "rb").read(),
                    file_name="generated_proposal.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with col2:
                st.download_button(
                    label="Download as PDF",
                    data=open(pdf_path, "rb").read(),
                    file_name="generated_proposal.pdf",
                    mime="application/pdf"
                )

            # Clean up temporary files
            os.unlink(md_path)
            os.unlink(docx_path)
            os.unlink(pdf_path)

        except Exception as e:
            st.error(f"Error preparing downloads: {str(e)}")
            # Clean up on error
            if os.path.exists(md_path):
                os.unlink(md_path)
            if os.path.exists(docx_path):
                os.unlink(docx_path)
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)

# Check if we should display the full proposal view
if "view_full_proposal" in st.session_state and st.session_state.view_full_proposal:
    st.title("Full Proposal Preview")
    
    # Add a button to go back to the main interface
    if st.button("← Back to Generator"):
        st.session_state.view_full_proposal = False
        st.experimental_rerun()
    
    # Display the complete proposal
    complete_proposal = get_complete_proposal()
    st.markdown(complete_proposal)
    
    # Add download options here too
    with tempfile.NamedTemporaryFile(delete=False, suffix=".md") as md_file:
        md_file.write(complete_proposal.encode())
        md_path = md_file.name

    docx_path = md_path.replace(".md", ".docx")
    pdf_path = md_path.replace(".md", ".pdf")

    try:
        # Convert to DOCX and PDF with extracted images
        convert_md_to_docx(complete_proposal, docx_path, st.session_state.extracted_images)
        convert_md_to_pdf(complete_proposal, pdf_path, st.session_state.extracted_images)

        # Create download buttons
        col1, col2 = st.columns(2)

        with col1:
            st.download_button(
                label="Download as DOCX",
                data=open(docx_path, "rb").read(),
                file_name="generated_proposal.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with col2:
            st.download_button(
                label="Download as PDF",
                data=open(pdf_path, "rb").read(),
                file_name="generated_proposal.pdf",
                mime="application/pdf"
            )

        # Clean up temporary files
        os.unlink(md_path)
        os.unlink(docx_path)
        os.unlink(pdf_path)

    except Exception as e:
        st.error(f"Error preparing downloads: {str(e)}")
        # Clean up on error
        if os.path.exists(md_path):
            os.unlink(md_path)
        if os.path.exists(docx_path):
            os.unlink(docx_path)
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)

# Add helpful information in the sidebar
st.sidebar.title("Help")
st.sidebar.info("""
### How to use this app:
1. Enter project requirements via text or file upload
2. Upload PDF/DOCX documents to build your knowledge base
3. Create the knowledge base using the button
4. Choose a proposal template (default, edit, or upload)
5. Click "Generate Next Section" to create your proposal section by section
6. Review and edit each section before generating the next
7. View the complete proposal in the preview area
8. Download in your preferred format (DOCX or PDF)
""")

st.sidebar.title("About")
st.sidebar.info("""
This app uses:
- LangChain with Hugging Face embeddings for the vector database
- Groq for fast, high-quality text generation
- Streamlit for the interactive interface
- PyMuPDF and python-docx for document processing and image extraction
- Mermaid for generating custom workflow diagrams

To get a Groq API key, sign up at https://console.groq.com/
""")

